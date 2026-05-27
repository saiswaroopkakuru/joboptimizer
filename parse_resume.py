import sys
import os
import re
import json
from docx import Document

def clean_text(text):
    return text.strip().replace('\xa0', ' ').replace('\u200b', '')

def parse_resume_data(docx_path):
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Resume file not found at: {docx_path}")
        
    doc = Document(docx_path)
    
    # 1. Extract contact details
    full_text = []
    paragraphs = []
    for p in doc.paragraphs:
        txt = clean_text(p.text)
        if txt:
            paragraphs.append(txt)
        full_text.append(txt)
    
    # Defaults
    profile = {
        "firstName": "Sai Swaroop",
        "lastName": "Reddy",
        "email": "saiswaroopkakuru@gmail.com",
        "phone": "+1 (408) 590-8917",
        "phoneClean": "4085908917",
        "linkedin": "https://www.linkedin.com/in/venkatasaiswaroopkakuru/",
        "github": "https://github.com/venkatasaiswaroopkakuru",
        "location": "San Jose, CA",
        "city": "San Jose",
        "state": "CA",
        "zip": "95112",
        "country": "United States",
        "experience": [],
        "education": [],
        "skills": []
    }
    
    # Scan text for email, phone, links
    joined_text = "\n".join(paragraphs)
    
    # Find email
    email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', joined_text)
    if email_match:
        profile["email"] = email_match.group(0)
        
    # Find phone
    phone_match = re.search(r'(?:\+?1[-. ]?)?\(?([0-9]{3})\)?[-. ]?([0-9]{3})[-. ]?([0-9]{4})', joined_text)
    if phone_match:
        profile["phone"] = f"+1 ({phone_match.group(1)}) {phone_match.group(2)}-{phone_match.group(3)}"
        profile["phoneClean"] = f"{phone_match.group(1)}{phone_match.group(2)}{phone_match.group(3)}"
        
    # Find LinkedIn
    li_match = re.search(r'(linkedin\.com/in/[a-zA-Z0-9_-]+)', joined_text)
    if li_match:
        profile["linkedin"] = "https://" + li_match.group(0)
        
    # Find GitHub
    gh_match = re.search(r'(github\.com/[a-zA-Z0-9_-]+)', joined_text)
    if gh_match:
        profile["github"] = "https://" + gh_match.group(0)

    # 2. Extract Work Experience
    # We parse the experience section block-by-block.
    experience_list = []
    
    current_company = None
    current_title = None
    current_dates = None
    current_bullets = []
    
    in_exp = False
    
    # Simple regex for date range (e.g. "Aug 2024 - Present" or "Jan 2023 - May 2024")
    date_range_regex = r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s*-\s*(?:Present|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4})'
    
    for p in doc.paragraphs:
        p_text = clean_text(p.text)
        p_text_lower = p_text.lower()
        
        # Section triggers
        if p_text_lower == "experience":
            in_exp = True
            continue
        if in_exp and (p_text_lower == "technical skills" or p_text_lower == "education & certifications" or p_text_lower.startswith("languages & concepts:")):
            in_exp = False
            if current_company:
                experience_list.append({
                    "company": current_company,
                    "title": current_title,
                    "dates": current_dates,
                    "bullets": current_bullets
                })
                current_company = None
            continue
            
        if in_exp and p_text:
            d_match = re.search(date_range_regex, p_text)
            if d_match:
                if current_company:
                    experience_list.append({
                        "company": current_company,
                        "title": current_title,
                        "dates": current_dates,
                        "bullets": current_bullets
                    })
                    current_bullets = []
                    current_title = None
                
                current_dates = d_match.group(0)
                # Split the rest to get company and location
                header_without_date = p_text.replace(current_dates, "").strip()
                # Split by tab or multiple spaces
                parts = re.split(r'\t|\s{2,}', header_without_date)
                company_loc = parts[0].strip()
                if " - " in company_loc:
                    comp_parts = company_loc.split(" - ")
                    current_company = comp_parts[0].strip()
                else:
                    current_company = company_loc
            elif current_company and not current_title:
                current_title = p_text
            elif current_company:
                bullet_clean = re.sub(r'^[\u2022\u25e6\u25aa\-*]\s*', '', p_text).strip()
                if bullet_clean:
                    current_bullets.append(bullet_clean)
                    
    if current_company:
        experience_list.append({
            "company": current_company,
            "title": current_title,
            "dates": current_dates,
            "bullets": current_bullets
        })
        
    profile["experience"] = experience_list

    # 3. Extract Education
    # Base resume contains:
    # "M.S. Computer Science - Southern Illinois University Carbondale	Jan 2023 - May 2024"
    education_list = []
    in_edu = False
    
    for p in doc.paragraphs:
        p_text = clean_text(p.text)
        p_text_lower = p_text.lower()
        
        if "education" in p_text_lower and not "experience" in p_text_lower:
            in_edu = True
            continue
        if in_edu and ("experience" in p_text_lower or "technical skills" in p_text_lower or "certifications" in p_text_lower or p_text.startswith("Languages & Concepts:")):
            in_edu = False
            continue
            
        if in_edu and p_text:
            if "southern illinois university" in p_text_lower:
                edu_entry = {
                    "school": "Southern Illinois University Carbondale",
                    "degree": "Master of Science",
                    "fieldOfStudy": "Computer Science",
                    "dates": "Jan 2023 - May 2024",
                    "gradYear": "2024"
                }
                education_list.append(edu_entry)
                
    if not education_list:
        # Fallback default
        education_list.append({
            "school": "Southern Illinois University Carbondale",
            "degree": "Master of Science",
            "fieldOfStudy": "Computer Science",
            "dates": "Jan 2023 - May 2024",
            "gradYear": "2024"
        })
        
    profile["education"] = education_list

    # 4. Extract Technical Skills
    skills_list = []
    prefixes = [
        "Languages & Concepts:",
        "Backend & APIs:",
        "Cloud & Infrastructure:",
        "Containers & CI/CD:",
        "Data & Messaging:",
        "Testing & Observability:",
        "Security & Practices:"
    ]
    
    for p in doc.paragraphs:
        p_text = clean_text(p.text)
        for prefix in prefixes:
            if p_text.startswith(prefix):
                content = p_text[len(prefix):].strip()
                # Split by | and , to get raw skills
                sections = content.split("|")
                for sec in sections:
                    items = sec.split(",")
                    for item in items:
                        # Clean item like "Python (3.8-3.12)" to "Python"
                        cleaned = re.sub(r'\(.*?\)', '', item).strip()
                        if cleaned and cleaned not in skills_list:
                            skills_list.append(cleaned)
                            
    profile["skills"] = skills_list
    
    return profile

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python parse_resume.py <path_to_docx>")
        sys.exit(1)
        
    path = sys.argv[1]
    try:
        data = parse_resume_data(path)
        print(json.dumps(data, indent=2))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)
