import os
import sys
import shutil
import re
import json
import urllib.request
import urllib.parse
import ssl
import random
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SKILL_CAPITALIZATION = {
    "python": "Python",
    "javascript": "JavaScript",
    "typescript": "TypeScript",
    "sql": "SQL",
    "pl/sql": "PL/SQL",
    "bash": "Bash",
    "yaml": "YAML",
    "groovy": "Groovy",
    "java": "Java",
    "c++": "C++",
    "c#": "C#",
    "go": "Go",
    "golang": "Go",
    "ruby": "Ruby",
    "rust": "Rust",
    "html": "HTML",
    "css": "CSS",
    "php": "PHP",
    "scala": "Scala",
    "kotlin": "Kotlin",
    "swift": "Swift",
    "oop": "OOP",
    "tdd": "TDD",
    "react": "React",
    "react.js": "React.js",
    "angular": "Angular",
    "vue": "Vue.js",
    "jquery": "jQuery",
    "bootstrap": "Bootstrap",
    "fastapi": "FastAPI",
    "fast api": "FastAPI",
    "django": "Django",
    "flask": "Flask",
    "sqlalchemy": "SQLAlchemy",
    "pydantic": "Pydantic",
    "celery": "Celery",
    "restful": "RESTful",
    "rest": "REST",
    "restful apis": "RESTful APIs",
    "graphql": "GraphQL",
    "grpc": "gRPC",
    "soap": "SOAP",
    "xml": "XML",
    "openapi": "OpenAPI",
    "swagger": "Swagger",
    "microservices": "Microservices",
    "aws": "AWS",
    "amazon web services": "Amazon Web Services",
    "gcp": "GCP",
    "google cloud": "Google Cloud",
    "azure": "Azure",
    "terraform": "Terraform",
    "aws cdk": "AWS CDK",
    "cdk": "CDK",
    "pulumi": "Pulumi",
    "cloudformation": "CloudFormation",
    "ansible": "Ansible",
    "iac": "IaC",
    "infrastructure as code": "Infrastructure as Code",
    "ec2": "Amazon EC2",
    "s3": "Amazon S3",
    "lambda": "AWS Lambda",
    "rds": "Amazon RDS",
    "eks": "Amazon EKS",
    "boto3": "Boto3",
    "step functions": "AWS Step Functions",
    "api gateway": "Amazon API Gateway",
    "docker": "Docker",
    "kubernetes": "Kubernetes",
    "k8s": "Kubernetes",
    "helm": "Helm",
    "istio": "Istio",
    "jenkins": "Jenkins",
    "github actions": "GitHub Actions",
    "gitlab ci/cd": "GitLab CI/CD",
    "argocd": "ArgoCD",
    "codepipeline": "AWS CodePipeline",
    "gitops": "GitOps",
    "ci/cd": "CI/CD",
    "postgresql": "PostgreSQL",
    "postgres": "PostgreSQL",
    "mysql": "MySQL",
    "mongodb": "MongoDB",
    "mongo": "MongoDB",
    "redis": "Redis",
    "oracle": "Oracle",
    "aurora": "Amazon Aurora",
    "kafka": "Apache Kafka",
    "apache kafka": "Apache Kafka",
    "rabbitmq": "RabbitMQ",
    "sqs": "Amazon SQS",
    "sns": "Amazon SNS",
    "caching": "Caching",
    "sql server": "SQL Server",
    "mssql": "MS SQL Server",
    "pytest": "pytest",
    "unittest": "unittest",
    "locust": "Locust",
    "prometheus": "Prometheus",
    "grafana": "Grafana",
    "elk": "ELK Stack",
    "elasticsearch": "Elasticsearch",
    "datadog": "Datadog",
    "splunk": "Splunk",
    "cloudwatch": "Amazon CloudWatch",
    "opentelemetry": "OpenTelemetry",
    "oauth2": "OAuth2",
    "jwt": "JWT",
    "sonarqube": "SonarQube",
    "snyk": "Snyk",
    "trivy": "Trivy",
    "agile": "Agile",
    "scrum": "Scrum",
    "jira": "Jira",
    "git": "Git",
    "linux": "Linux",
    "unix": "Unix",
    
    # New additions
    "node": "Node.js",
    "node.js": "Node.js",
    "nodejs": "Node.js",
    "api": "API",
    "apis": "APIs",
    "restful api": "RESTful API",
    "rest api": "REST API",
    "soap/xml": "SOAP/XML",
    "openapi/swagger": "OpenAPI/Swagger",
    "django rest framework": "Django REST Framework",
    "unittest.mock": "unittest.mock",
    "pylint": "pylint",
    "mypy": "mypy",
    "pyodbc": "pyodbc",
    "pymongo": "PyMongo",
    "motor": "Motor",
    "confluent-kafka-python": "confluent-kafka-python",
}

def capitalize_skill(s):
    s_clean = s.strip()
    if not s_clean:
        return ""
    s_lower = s_clean.lower()
    
    if s_lower in SKILL_CAPITALIZATION:
        return SKILL_CAPITALIZATION[s_lower]
        
    # If the original string `s` has some uppercase letters, we should probably keep it as is.
    if s_clean != s_lower:
        return s_clean
        
    # If it is all lowercase, use fallback rules:
    if len(s_lower) <= 4:
        lowercase_short_words = {"java", "go", "rust", "helm", "git", "unix", "bash", "perl", "chef", "ruby"}
        if s_lower in lowercase_short_words:
            return s_clean.title()
        return s_clean.upper()
        
    return s_clean.title()

# Mapping of common skills to their corresponding resume sections
CATEGORIES_MAPPING = {
    # Languages & Concepts
    "python": "languages", "javascript": "languages", "typescript": "languages", "sql": "languages",
    "pl/sql": "languages", "bash": "languages", "yaml": "languages", "groovy": "languages",
    "java": "languages", "c++": "languages", "c#": "languages", "go": "languages", "golang": "languages",
    "ruby": "languages", "rust": "languages", "html": "languages", "css": "languages", "php": "languages",
    "scala": "languages", "kotlin": "languages", "swift": "languages", "oop": "languages",
    "algorithms": "languages", "data structures": "languages", "tdd": "languages",
    "react": "languages", "react.js": "languages", "angular": "languages", "vue": "languages",
    "jquery": "languages", "bootstrap": "languages",

    # Backend & APIs
    "fastapi": "backend", "fast api": "backend", "django": "backend", "flask": "backend",
    "sqlalchemy": "backend", "pydantic": "backend", "celery": "backend", "restful": "backend",
    "rest": "backend", "restful apis": "backend", "graphql": "backend", "grpc": "backend",
    "soap": "backend", "xml": "backend", "openapi": "backend", "swagger": "backend",
    "microservices": "backend", "event-driven": "backend", "distributed systems": "backend",
    "spring": "backend", "spring boot": "backend", "node.js": "backend", "node": "backend", "express": "backend",

    # Cloud & Infrastructure
    "aws": "cloud", "amazon web services": "cloud", "gcp": "cloud", "google cloud": "cloud",
    "azure": "cloud", "terraform": "cloud", "aws cdk": "cloud", "cdk": "cloud", "pulumi": "cloud",
    "cloudformation": "cloud", "ansible": "cloud", "iac": "cloud", "infrastructure as code": "cloud",
    "ec2": "cloud", "s3": "cloud", "lambda": "cloud", "rds": "cloud", "eks": "cloud", "boto3": "cloud",
    "step functions": "cloud", "api gateway": "cloud",

    # Containers & CI/CD
    "docker": "containers", "kubernetes": "containers", "k8s": "containers", "helm": "containers",
    "istio": "containers", "jenkins": "containers", "github actions": "containers",
    "gitlab ci/cd": "containers", "argocd": "containers", "codepipeline": "containers",
    "gitops": "containers", "ci/cd": "containers",

    # Data & Messaging
    "postgresql": "data", "postgres": "data", "mysql": "data", "mongodb": "data", "mongo": "data",
    "redis": "data", "oracle": "data", "aurora": "data", "kafka": "data", "apache kafka": "data",
    "rabbitmq": "data", "sqs": "data", "sns": "data", "event streaming": "data", "caching": "data",
    "sql server": "data", "mssql": "data",

    # Testing & Observability
    "pytest": "testing", "unittest": "testing", "locust": "testing", "prometheus": "testing",
    "grafana": "testing", "elk": "testing", "elasticsearch": "testing", "datadog": "testing",
    "splunk": "testing", "cloudwatch": "testing", "opentelemetry": "testing",

    # Security & Practices
    "oauth2": "security", "jwt": "security", "auth": "security", "authentication": "security",
    "authorization": "security", "sonarqube": "security", "snyk": "security", "trivy": "security",
    "agile": "security", "scrum": "security", "jira": "security", "git": "security", "linux": "security",
    "unix": "security"
}

PREFIX_MAP = {
    "languages": "Languages & Concepts:",
    "backend": "Backend & APIs:",
    "cloud": "Cloud & Infrastructure:",
    "containers": "Containers & CI/CD:",
    "data": "Data & Messaging:",
    "testing": "Testing & Observability:",
    "security": "Security & Practices:"
}

def get_category(skill):
    sk_lower = skill.lower().strip()
    if sk_lower in CATEGORIES_MAPPING:
        return CATEGORIES_MAPPING[sk_lower]
    if any(x in sk_lower for x in ["cloud", "aws", "gcp", "azure", "infra", "boto"]):
        return "cloud"
    if any(x in sk_lower for x in ["docker", "k8s", "kube", "ci", "cd", "pipeline", "gitops", "jenkins", "action"]):
        return "containers"
    if any(x in sk_lower for x in ["test", "observ", "alert", "log", "monitor", "prometheus", "grafana", "splunk", "datadog"]):
        return "testing"
    if any(x in sk_lower for x in ["db", "sql", "cache", "kafka", "mq", "stream", "redis", "postgres", "mongo"]):
        return "data"
    if any(x in sk_lower for x in ["api", "rest", "microservice", "framework", "fastapi", "django", "flask", "spring"]):
        return "backend"
    if any(x in sk_lower for x in ["auth", "security", "jwt", "oauth2", "agile", "scrum", "jira"]):
        return "security"
    return "languages"

def is_skill_present(paragraph_text, skill):
    skill_clean = skill.lower().strip()
    text_clean = paragraph_text.lower()
    
    patterns = [skill_clean]
    if skill_clean == "fastapi":
        patterns.append("fast api")
    elif skill_clean == "fast api":
        patterns.append("fastapi")
        
    if " " in skill_clean:
        patterns.append(skill_clean.replace(" ", ""))
        
    for pat in patterns:
        escaped_pat = re.escape(pat)
        regex_pattern = f"(?<![a-z0-9+#]){escaped_pat}(?![a-z0-9+#])"
        if re.search(regex_pattern, text_clean):
            return True
            
    return False

def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = new_text

def update_prefixed_paragraph(p, prefix, body_text):
    prefix_clean = prefix.strip()
    body_clean = body_text.strip()
    
    font_name = None
    font_size = None
    if p.runs:
        font_name = p.runs[0].font.name
        font_size = p.runs[0].font.size
        
    p.text = ""
    
    # Prefix run (bold)
    r_prefix = p.add_run(prefix_clean)
    r_prefix.bold = True
    if font_name:
        r_prefix.font.name = font_name
    if font_size:
        r_prefix.font.size = font_size
        
    # Body run (not bold)
    r_body = p.add_run(" " + body_clean)
    r_body.bold = False
    if font_name:
        r_body.font.name = font_name
    if font_size:
        r_body.font.size = font_size

def get_bullet_skills(lst, count, fallback_str):
    if not lst:
        return fallback_str
    capitalized = []
    for s in lst:
        capitalized.append(capitalize_skill(s))
    seen_cap = set()
    unique = []
    for cap in capitalized:
        if cap not in seen_cap:
            seen_cap.add(cap)
            unique.append(cap)
    return ", ".join(unique[:count])

def generate_optimized_bullets(company, title, matched, missing, selected_llm="claude"):
    # Put missing skills first, then matched skills (to prioritize missing skills in optimization)
    seen = set()
    all_skills = []
    for s in missing + matched:
        s_lower = s.lower().strip()
        if s_lower and s_lower not in seen:
            seen.add(s_lower)
            all_skills.append(s_lower)
    
    def filter_skills(keywords_list):
        return [s for s in all_skills if s in keywords_list]
        
    job_languages = filter_skills(["python", "javascript", "typescript", "java", "go", "golang", "c++", "c#", "ruby", "rust", "sql", "bash"])
    job_backend = filter_skills(["fastapi", "django", "flask", "sqlalchemy", "spring", "spring boot", "node.js", "node", "express", "graphql", "grpc", "restful", "rest", "celery"])
    job_cloud = filter_skills(["aws", "gcp", "azure", "terraform", "cdk", "pulumi", "cloudformation", "ansible", "ec2", "s3", "lambda", "rds", "boto3", "step functions", "api gateway"])
    job_containers = filter_skills(["kubernetes", "docker", "helm", "eks", "aks", "gke", "argocd", "gitops", "jenkins", "ci/cd"])
    job_data = filter_skills(["postgresql", "mysql", "mongodb", "redis", "oracle", "kafka", "rabbitmq", "sqs", "sns"])
    job_obs = filter_skills(["splunk", "prometheus", "grafana", "datadog", "cloudwatch", "opentelemetry", "elk"])

    title_lower = title.lower()
    
    # Determine job type and defaults
    if any(x in title_lower for x in ["sre", "devops", "infrastructure", "platform", "production"]):
        job_type = "sre"
        default_lang = "Python and Bash"
        default_backend = "gRPC and REST"
        default_cloud = "AWS and Terraform"
        default_containers = "Kubernetes and Docker"
        default_data = "Apache Kafka"
        default_obs = "Prometheus and Grafana"
        default_api = "gRPC and REST APIs"
        default_security = "OAuth2/JWT"
    elif any(x in title_lower for x in ["frontend", "ui", "react", "web", "fullstack", "full stack"]):
        job_type = "frontend"
        default_lang = "TypeScript and JavaScript"
        default_backend = "React"
        default_cloud = "AWS"
        default_containers = "Docker and GitHub Actions"
        default_data = "RESTful APIs"
        default_obs = "Datadog"
        default_api = "REST and GraphQL"
        default_security = "JWT"
    elif any(x in title_lower for x in ["java", "spring"]):
        job_type = "java"
        default_lang = "Java and SQL"
        default_backend = "Spring Boot"
        default_cloud = "AWS and Terraform"
        default_containers = "Kubernetes and Docker"
        default_data = "Apache Kafka and PostgreSQL"
        default_obs = "Prometheus and Grafana"
        default_api = "Spring REST APIs"
        default_security = "Spring Security (OAuth2/JWT)"
    else:
        job_type = "backend"
        default_lang = "Python"
        default_backend = "FastAPI and Django"
        default_cloud = "AWS and Terraform"
        default_containers = "Kubernetes and Docker"
        default_data = "Apache Kafka and PostgreSQL"
        default_obs = "Splunk, Prometheus, and Grafana"
        default_api = "FastAPI REST APIs"
        default_security = "OAuth2/JWT"

    lang_str = get_bullet_skills(job_languages, 2, default_lang)
    backend_str = get_bullet_skills(job_backend, 2, default_backend)
    cloud_str = get_bullet_skills(job_cloud, 2, default_cloud)
    container_str = get_bullet_skills(job_containers, 2, default_containers)
    data_str = get_bullet_skills(job_data, 2, default_data)
    obs_str = get_bullet_skills(job_obs, 2, default_obs)
    
    # Helpers for specialized fields
    api_str = get_bullet_skills(filter_skills(["rest", "restful", "restful apis", "graphql", "grpc", "soap", "xml", "openapi", "swagger"]), 2, default_api)
    security_str = get_bullet_skills(filter_skills(["oauth2", "jwt", "auth", "authentication", "authorization", "spring security", "security"]), 2, default_security)
    cicd_str = get_bullet_skills(filter_skills(["jenkins", "github actions", "gitlab ci/cd", "argocd", "codepipeline", "ci/cd"]), 3, "Jenkins, GitHub Actions, and ArgoCD")

    templates = []
    
    if job_type == "sre":
        if selected_llm == "claude":
            templates = [
                f"Spearheaded platform reliability and scaling initiatives for HSBC's infrastructure, leveraging {cloud_str} to automate environment provisioning and reduce manual setup overhead by 90% across 15+ microservices.",
                f"Orchestrated highly available containerized deployments using {container_str}, configuring custom autoscaling policies and service meshes to handle peak traffic surges of 1M+ daily transactions.",
                f"Modernized CI/CD pipelines using {cicd_str} with GitOps workflows; accelerated release cycles by 20% and enabled blue/green and canary deployments of fraud-detection rules with automated rollback.",
                f"Integrated {data_str} for event streaming and distributed data platforms; optimized multi-stage container builds (image size -25%) and configured high-performance routing to absorb peak network traffic.",
                f"Implemented comprehensive end-to-end monitoring and alerting networks with {obs_str}, improving cluster observability and reducing incident MTTR by 20% for HSBC operations."
            ]
        elif selected_llm == "opus":
            templates = [
                f"Formulated architectural frameworks for HSBC's global SRE divisions, deploying scalable {cloud_str} cloud resources to eliminate 90% of manual deployment workloads.",
                f"Engineered highly resilient, zero-trust container configurations using {container_str}, orchestrating multi-region node clusters to sustain peak traffic of 1M+ transactions.",
                f"Architected sophisticated continuous integration pathways via {cicd_str}, implementing declarative GitOps pipelines and automated rollbacks for system stability.",
                f"Spearheaded low-latency streaming patterns by coupling {data_str} with container orchestration, decreasing deployment footprints by 25% via multi-stage packaging.",
                f"Constructed unified metrics and alert meshes using {obs_str}, creating comprehensive observability dashboards to reduce operational MTTR by 20%."
            ]
        elif selected_llm == "gemini":
            templates = [
                f"Pioneered infrastructure scaling and cloud reliability initiatives for HSBC, leveraging {cloud_str} to build self-healing deployment environments and reduce setup time by 90%.",
                f"Designed and deployed highly resilient containerized architectures using {container_str}, establishing automated load balancing to manage 1M+ daily transaction spikes.",
                f"Partnered with development teams to roll out modern CI/CD pipelines with {cicd_str}, enabling seamless, low-risk canary deployments and automated rollbacks.",
                f"Leveraged {data_str} for high-performance event-driven data streaming; reduced container footprint by 25% through optimized configurations and custom scaling policies.",
                f"Enhanced overall production observability by implementing centralized monitoring via {obs_str}, empowering teams to trace requests and resolve incidents 20% faster."
            ]
        elif selected_llm == "gemini35":
            templates = [
                f"Pioneered enterprise-wide infrastructure automation at HSBC, employing {cloud_str} to model self-healing platform environments and cut environment setup time by 90%.",
                f"Designed and deployed responsive container topologies using {container_str}, configuring dynamic service meshes and HPA autoscaling to resolve 1M+ transaction spikes.",
                f"Facilitated seamless pipeline deliveries using {cicd_str} under GitOps models, speeding up automated deployments by 20% with safe canary testing.",
                f"Structured high-speed data flow architectures via {data_str}, utilizing container multi-stage optimizations to cut image footprints by 25%.",
                f"Maximized system transparency by implementing telemetry monitoring networks via {obs_str}, enabling operational engineers to isolate incidents 20% faster."
            ]
        elif selected_llm == "deepseek":
            templates = [
                f"Analyzed and refactored HSBC's cloud resource allocation, using {cloud_str} to automate infrastructure provisioning and eliminate 90% of manual configuration bottlenecks.",
                f"Optimized container orchestration using {container_str}, tuning resource limits and autoscaling parameters to safely process 1M+ daily transaction events.",
                f"Streamlined CI/CD delivery workflows utilizing {cicd_str} and GitOps, decreasing release cycle overhead by 20% while safeguarding production via automated rollbacks.",
                f"Resolved performance bottlenecks in event-driven streaming by integrating {data_str}, utilizing optimized multi-stage container builds to reduce footprint by 25%.",
                f"Constructed robust system observability models with {obs_str}, establishing telemetry dashboards that reduced mean time to resolution (MTTR) by 20%."
            ]
        elif selected_llm == "llama":
            templates = [
                f"Drove the transition to automated infrastructure at HSBC, using {cloud_str} to provision cloud resources and cut manual setup times by 90% across 15+ services.",
                f"Built and managed scalable container networks using {container_str}, ensuring high availability and seamless load distribution for 1M+ daily transactions.",
                f"Automated build and release operations via custom {cicd_str} pipelines, accelerating sprint delivery cycles by 20% with safe canary rollout strategies.",
                f"Integrated {data_str} to power event-driven microservices, applying best practices in container builds to shrink image sizes by 25% and improve startup times.",
                f"Spun up comprehensive monitoring and alerting systems using {obs_str}, providing deep visibility into application health and reducing service downtime by 20%."
            ]
        elif selected_llm == "gpt5":
            templates = [
                f"Architected and standardized cloud platform provisioning for HSBC systems, utilizing {cloud_str} to achieve 90% automation in infrastructure setup.",
                f"Orchestrated container cluster provisioning with {container_str}, designing robust ingress control and network policies to process 1M+ transactions daily.",
                f"Designed secure, high-concurrency CI/CD platforms using {cicd_str}, increasing automated delivery frequency by 20% with automated canary rollbacks.",
                f"Integrated distributed message streaming with {data_str}, achieving 25% storage optimization across container infrastructure via multi-stage builds.",
                f"Engineered centralized system telemetry architectures with {obs_str}, reducing mean time to detection and resolution (MTTR) by 20% for SRE divisions."
            ]
        elif selected_llm == "gemma":
            templates = [
                f"Drove open-source platform configurations at HSBC, using {cloud_str} to automate environment builds and eliminate 90% of manual sysadmin overhead.",
                f"Deployed lightweight container clusters with {container_str}, managing load-balancers and resource specifications to process 1M+ transactions.",
                f"Configured developer-friendly CI/CD and deployment tasks using {cicd_str}, helping accelerate delivery timelines by 20% with automated rollbacks.",
                f"Managed event streaming queues with {data_str}, optimizing container builds to achieve 25% smaller image sizes and faster startup speeds.",
                f"Integrated open monitoring stacks with {obs_str} for cluster telemetry, establishing dashboards that cut incident troubleshooting times by 20%."
            ]
        else: # gpt4o / default
            templates = [
                f"Managed cloud infrastructure provisioning and reliability for HSBC, utilizing {cloud_str} to create automated environment setups and reduce manual overhead by 90%.",
                f"Implemented scalable container deployments using {container_str}, configuring horizontal autoscaling to handle peak transaction volumes of 1M+ daily.",
                f"Configured stable CI/CD pipelines using {cicd_str} for automated deployments; improved release efficiency by 20% with canary testing mechanisms.",
                f"Integrated event-driven data streaming components using {data_str}; optimized multi-stage container builds to achieve 25% size reduction and faster startup.",
                f"Set up standard monitoring and alerting solutions using {obs_str}, improving application tracking metrics and reducing incident resolution times by 20%."
            ]
    elif job_type == "frontend":
        if selected_llm == "claude":
            templates = [
                f"Designed and developed responsive, high-performance user interfaces for HSBC's platform, utilizing {lang_str} and modern frontend frameworks like {backend_str} to build component-driven architectures.",
                f"Optimized web application performance, reducing bundle sizes by 30% and implementing caching structures to guarantee fluid transitions and sub-second load times on {cloud_str} cloud infrastructure.",
                f"Modernized CI/CD and deployment pipelines using {cicd_str} with automated testing and GitOps workflows; accelerated release cycles by 20% and enabled seamless canary rollouts.",
                f"Integrated RESTful APIs and distributed state management to process event streaming of 1M+ daily transaction events; optimized build assets and configured responsive frontend structures.",
                f"Designed UI patterns and telemetry logging using {api_str}; integrated {security_str} authentication and end-to-end front-end monitoring with {obs_str} to reduce incident MTTR."
            ]
        elif selected_llm == "opus":
            templates = [
                f"Formulated modern interface structures for HSBC's core platforms using {lang_str} and {backend_str}, focusing on component reuse and sub-second load metrics.",
                f"Optimized user-interaction latency and client-side rendering pathways, dropping bundle sizes by 30% across global {cloud_str} distributions.",
                f"Architected end-to-end integration pathways for API state management using {cicd_str}, enabling automated validation and canary interface rollouts.",
                f"Integrated telemetry event streams and state engines using {api_str} to handle 1M+ transaction views without UI main-thread blockage.",
                f"Designed security login workflows with {security_str} and comprehensive telemetry instrumentation via {obs_str} to lower frontend MTTR by 20%."
            ]
        elif selected_llm == "gemini":
            templates = [
                f"Pioneered interactive and responsive user experiences for HSBC's platform, leveraging {lang_str} and {backend_str} to create fluid, accessible frontend designs.",
                f"Boosted frontend speed and bundle efficiency by 30% on {cloud_str}, implementing advanced caching and progressive loading for a seamless user journey.",
                f"Built automated test suites and CI/CD workflows using {cicd_str}, ensuring high-quality, continuous interface updates with 20% faster deployment times.",
                f"Connected frontend states with {api_str} endpoints to handle real-time data flows; worked on optimizing user flows for 1M+ transaction views.",
                f"Implemented user analytics and error tracking with {obs_str}, securing access points using {security_str} and improving issue resolution speed by 20%."
            ]
        elif selected_llm == "gemini35":
            templates = [
                f"Pioneered highly interactive and accessible web interfaces for HSBC using {lang_str} and {backend_str}, ensuring high-quality component design.",
                f"Optimized client rendering speeds and asset caching on {cloud_str}, reducing bundle overhead by 30% to maximize screen loading efficiency.",
                f"Coordinated automated quality pipelines using {cicd_str}, running unit and visual validation tests to accelerate UI releases by 20%.",
                f"Connected client interfaces with real-time API feeds using {api_str}, rendering complex transaction views for 1M+ daily telemetry reports.",
                f"Implemented user-interaction flow analysis using {obs_str} and secured authentication interfaces using {security_str} to cut error times by 20%."
            ]
        elif selected_llm == "deepseek":
            templates = [
                f"Refactored HSBC's frontend codebase to use {lang_str} and {backend_str}, modularizing UI components to enhance rendering performance and developer velocity.",
                f"Optimized client-side rendering bottlenecks, reducing bundle sizes by 30% and deploying assets efficiently via {cloud_str} cloud networks.",
                f"Integrated automated unit and end-to-end testing with {cicd_str} pipelines, reducing regression issues and shortening release cycles by 20%.",
                f"Mapped complex data schemas to frontend states using {api_str}, handling real-time telemetry inputs of 1M+ transactions without lag.",
                f"Constructed secure client-side authentication guards with {security_str} and telemetry logging using {obs_str}, dropping frontend exception MTTR by 20%."
            ]
        elif selected_llm == "llama":
            templates = [
                f"Built beautiful, high-fidelity web views for HSBC, writing clean {lang_str} code using {backend_str} to construct reusable component systems.",
                f"Drove bundle size reduction of 30% and set up asset caching to ensure fast UI response times, hosting assets on {cloud_str} for global scale.",
                f"Automated client deployments via {cicd_str} workflows, enabling rapid UI changes and shortening validation times by 20%.",
                f"Consumed high-throughput REST APIs and data streams from {api_str}, building responsive UI components to display real-time transaction updates.",
                f"Integrated secure login flows via {security_str} and set up monitoring with {obs_str} to catch UI errors before they hit production."
            ]
        elif selected_llm == "gpt5":
            templates = [
                f"Architected component-driven frontend frameworks for HSBC utilizing {lang_str} and {backend_str} to establish scalable web design patterns.",
                f"Implemented advanced client-side performance optimizations on {cloud_str}, reducing asset bundle sizes by 30% and optimizing paint times.",
                f"Designed scalable test integration and deployment pipelines using {cicd_str}, achieving 20% faster turnaround on UI modifications.",
                f"Integrated high-throughput data models with frontend state engines using {api_str} to process 1M+ daily transaction dashboards.",
                f"Configured secure authentication guards via {security_str} and structured logging endpoints with {obs_str} for 20% faster issue tracking."
            ]
        elif selected_llm == "gemma":
            templates = [
                f"Developed accessible and modular client components using {lang_str} and {backend_str} to support web operations for HSBC.",
                f"Drove asset size reduction of 30% and configured global delivery networks on {cloud_str} to enable fast client page loads.",
                f"Automated interface builds and unit-test executions via {cicd_str}, reducing deployment errors and release cycle times by 20%.",
                f"Consumed structured backend services using {api_str} to present real-time dashboards for 1M+ transaction list views.",
                f"Integrated user authentication checks via {security_str} and client telemetry tracking with {obs_str} to speed up UI issue detection."
            ]
        else: # gpt4o / default
            templates = [
                f"Developed frontend user interfaces for HSBC applications, utilizing {lang_str} and {backend_str} to implement responsive design patterns.",
                f"Improved web performance metrics, reducing bundle sizes by 30% and configuring cloud hosting delivery networks on {cloud_str}.",
                f"Maintained automated UI test integration and release management with {cicd_str}, achieving 20% faster cycle runs for interface features.",
                f"Integrated API data sources using {api_str} to present real-time dashboards for 1M+ transaction records.",
                f"Implemented secure authentication protocols using {security_str} and frontend error tracking using {obs_str} for rapid troubleshooting."
            ]
    else: # backend, java, etc.
        if selected_llm == "claude":
            templates = [
                f"Engineered high-throughput backend services in {lang_str} using the {backend_str} framework for HSBC, processing business-critical workflows and optimizing database queries for 25% faster response times.",
                f"Architected and deployed cloud-native applications on {cloud_str} cloud infrastructure, managing container deployment pipelines using {container_str} to achieve seamless, zero-downtime rollouts.",
                f"Modernized CI/CD pipelines using {cicd_str} with GitOps workflows; accelerated release cycles by 20% and enabled blue/green and canary deployments of fraud-detection rules with automated rollback.",
                f"Integrated {data_str} for event streaming of 1M+ daily transaction events; optimized multi-stage {container_str if 'docker' in container_str.lower() or 'kubernetes' in container_str.lower() else 'Docker'} builds (image size -25%) and configured Kubernetes HPA autoscaling to absorb peak fraud-attack traffic.",
                f"Designed RESTful APIs and interfaces using {api_str} with comprehensive documentation; integrated {security_str} authentication and end-to-end observability with {obs_str} - including real-time anomaly alerts on fraud signals, reducing incident MTTR."
            ]
        elif selected_llm == "gemini":
            templates = [
                f"Pioneered core backend logic updates for HSBC's transaction engine in {lang_str} using {backend_str}, raising query speeds by 25% through indexing and caching.",
                f"Designed and deployed cloud microservices on {cloud_str}, configuring automated container provisioning with {container_str} for maximum uptime and reliability.",
                f"Collaborated on automating release operations using {cicd_str}, enabling safe canary rollouts and automated rollback strategies that cut delivery times by 20%.",
                f"Built robust data streaming pipes using {data_str} to process 1M+ transactions; trimmed container image footprints by 25% via multi-stage builds.",
                f"Created documented RESTful APIs with {api_str}, securing access points using {security_str} and introducing Prometheus monitoring via {obs_str} to trace issues."
            ]
        elif selected_llm == "opus":
            templates = [
                f"Formulated architectural frameworks for HSBC's global backend platforms using {lang_str} and {backend_str}, optimizing database query paths to increase response velocity by 25%.",
                f"Architected resilient, cloud-native backend layers on {cloud_str}, packaging distributed microservices using {container_str} to establish scalable node clusters.",
                f"Implemented sophisticated CI/CD release architectures via {cicd_str}, deploying declarative pipelines with GitOps models to eliminate release bottlenecks.",
                f"Spearheaded low-latency message streaming pipelines utilizing {data_str} to process 1M+ transactions, reducing container footprints by 25% via multi-stage builds.",
                f"Constructed zero-trust security integrations with {security_str} and telemetry logging using {obs_str}, reducing incident MTTR by 20% across API layers."
            ]
        elif selected_llm == "gemini35":
            templates = [
                f"Pioneered enterprise-wide backend service enhancements at HSBC using {lang_str} and {backend_str}, raising query speeds by 25% with dynamic index configurations.",
                f"Designed and deployed responsive backend topologies on {cloud_str}, configuring automated containers using {container_str} to handle transaction loads.",
                f"Coordinated automated delivery pipelines using {cicd_str} under GitOps models, accelerating releases by 20% with safe canary testing.",
                f"Structured high-speed data flow architectures via {data_str} to process 1M+ transactions, utilizing container multi-stage optimization to cut footprints by 25%.",
                f"Exposed documented APIs using {api_str}, securing access points via {security_str} and implementing telemetry monitoring networks via {obs_str}."
            ]
        elif selected_llm == "deepseek":
            templates = [
                f"Refactored and optimized high-performance backend systems in {lang_str} using {backend_str} for HSBC, analyzing query execution plans to speed up database response by 25%.",
                f"Designed cloud-native architectures on {cloud_str}, orchestrating container environments with {container_str} to ensure scalable service execution.",
                f"Constructed automated CI/CD and deployment workflows using {cicd_str}, ensuring high reliability of deployments and accelerating cycle speed by 20%.",
                f"Resolved throughput limits in data event processors by integrating {data_str}, achieving 25% container size reduction through optimized builds.",
                f"Developed secure API gateways using {api_str} and {security_str}, linking systems to telemetry nodes with {obs_str} to decrease incident MTTR by 20%."
            ]
        elif selected_llm == "llama":
            templates = [
                f"Built scale-optimized backend services in {lang_str} using {backend_str} at HSBC, scaling databases and APIs to improve transaction processing speed by 25%.",
                f"Provisioned and managed cloud infrastructure using {cloud_str}, packaging services in container clusters via {container_str} for rapid deployment.",
                f"Automated build and deploy operations using custom {cicd_str} pipelines, helping shorten sprint cycles by 20% with safe canary releases.",
                f"Leveraged {data_str} for high-performance event-driven data streaming of 1M+ daily events, applying multi-stage container builds to save 25% storage space.",
                f"Designed and exposed clean REST APIs using {api_str}; secured endpoints via {security_str} and set up monitoring with {obs_str} to trace errors."
            ]
        elif selected_llm == "gpt5":
            templates = [
                f"Architected and standardized backend systems for HSBC using {lang_str} and {backend_str}, optimizing database schemas to accelerate response times by 25%.",
                f"Orchestrated containerized microservices deployments on {cloud_str} cloud infrastructure using {container_str} to achieve 90% automation in platform provisioning.",
                f"Designed secure, high-concurrency CI/CD platforms using {cicd_str}, increasing automated delivery frequency by 20% with automated canary rollbacks.",
                f"Integrated distributed message streaming with {data_str} to process 1M+ daily transaction events, achieving 25% storage optimization across container infrastructure via multi-stage builds.",
                f"Engineered centralized system telemetry architectures with {obs_str} and API gateways using {api_str} and {security_str}, reducing MTTR by 20% for backend divisions."
            ]
        elif selected_llm == "gemma":
            templates = [
                f"Developed open-source backend configurations at HSBC using {lang_str} and {backend_str}, accelerating query processing and backend logic to reduce response latency by 25%.",
                f"Deployed lightweight container clusters with {container_str} on {cloud_str} cloud environments, managing container scaling and resource specifications.",
                f"Configured developer-friendly CI/CD and deployment tasks using {cicd_str}, helping accelerate delivery timelines by 20% with automated rollbacks.",
                f"Managed event streaming queues with {data_str} to process 1M+ transaction entries, optimizing container builds to achieve 25% smaller footprints.",
                f"Integrated open monitoring stacks with {obs_str} for cluster telemetry and exposed REST APIs with {api_str} secured using {security_str} to speed up issue detection."
            ]
        else: # gpt4o / default
            templates = [
                f"Developed backend services using {lang_str} and {backend_str} for HSBC systems, processing high volumes of transactions and optimizing database queries for 25% speedups.",
                f"Deployed cloud-native applications to {cloud_str} cloud platforms, managing container build pipelines using {container_str} for deployments.",
                f"Configured automated CI/CD pipelines using {cicd_str} to improve delivery efficiency and accelerate release deployment speeds by 20%.",
                f"Integrated messaging and database layers using {data_str} to process 1M+ transaction entries; optimized container sizes by 25% with multi-stage builds.",
                f"Designed RESTful APIs using {api_str} with OpenAPI documentation; integrated {security_str} authentication and monitoring utilities using {obs_str}."
            ]

    h = sum(ord(c) for c in company)
    metrics = [
        {"pct": "30%", "val": "1.2M", "time": "15%"},
        {"pct": "25%", "val": "950K", "time": "20%"},
        {"pct": "35%", "val": "2M", "time": "25%"},
        {"pct": "40%", "val": "1.5M", "time": "18%"}
    ]
    metric = metrics[h % len(metrics)]

    final_bullets = []
    for b in templates:
        b_mod = b.replace("30%", metric["pct"]).replace("25%", metric["pct"]).replace("20%", metric["time"]).replace("1M+", metric["val"] + "+")
        final_bullets.append(b_mod)

    return {
        "bullets": final_bullets,
        "lang_str": lang_str,
        "backend_str": backend_str,
        "cloud_str": cloud_str,
        "container_str": container_str,
        "data_str": data_str,
        "obs_str": obs_str,
        "api_str": api_str,
        "security_str": security_str,
        "cicd_str": cicd_str,
        "first_lang": get_bullet_skills(job_languages, 1, default_lang.split(" ")[0])
    }

# Sai's base skills for matched/missing classification
SAI_BASE_SKILLS = {
    "python", "javascript", "typescript", "sql", "pl/sql", "bash", "yaml", "groovy", "oop", "tdd",
    "fastapi", "fast api", "django", "flask", "sqlalchemy", "pydantic", "celery", "restful", "rest",
    "graphql", "grpc", "soap", "xml", "openapi", "swagger", "microservices", "aws", "gcp", "azure",
    "terraform", "aws cdk", "cdk", "pulumi", "cloudformation", "ansible", "iac", "ec2", "s3",
    "lambda", "rds", "eks", "boto3", "step functions", "api gateway", "docker", "kubernetes", "k8s",
    "helm", "istio", "jenkins", "github actions", "gitlab ci/cd", "argocd", "codepipeline", "gitops",
    "ci/cd", "postgresql", "mysql", "mongodb", "redis", "oracle", "aurora", "kafka", "apache kafka",
    "rabbitmq", "sqs", "sns", "pytest", "unittest", "locust", "prometheus", "grafana", "elk",
    "elasticsearch", "datadog", "splunk", "cloudwatch", "opentelemetry", "oauth2", "jwt",
    "sonarqube", "snyk", "trivy", "agile", "scrum", "jira", "git", "linux", "unix"
}

def search_ddg_tech_stack(company, title):
    query = f"{company} {title} careers tech stack"
    encoded_query = urllib.parse.quote(query)
    url = f"https://html.duckduckgo.com/html/?q={encoded_query}"
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }
    req = urllib.request.Request(url, headers=headers)
    context = ssl._create_unverified_context()
    
    found_skills = []
    try:
        with urllib.request.urlopen(req, context=context, timeout=8) as response:
            html = response.read().decode("utf-8", errors="ignore")
            
        with open("ddg_search_results.html", "w", encoding="utf-8") as f:
            f.write(html)
            
        # Extract snippets
        snippets = re.findall(r'<a class="result__snippet[^"]*"[^>]*>(.*?)</a>', html, re.DOTALL)
        combined_text = " ".join([re.sub(r'<[^>]*>', '', s) for s in snippets]).lower()
        
        for skill in SKILL_CAPITALIZATION.keys():
            escaped_pat = re.escape(skill)
            regex_pattern = f"(?<![a-z0-9+#]){escaped_pat}(?![a-z0-9+#])"
            if re.search(regex_pattern, combined_text):
                found_skills.append(skill)
    except Exception as e:
        print(f"Error scraping DuckDuckGo: {e}")
        
    return list(set(found_skills))

def get_job_description_from_db(company, title):
    try:
        if os.path.exists("jobs.json"):
            with open("jobs.json", "r", encoding="utf-8") as f:
                jobs = json.load(f)
            
            company_clean = company.lower().strip()
            title_clean = title.lower().strip()
            
            for job in jobs:
                j_comp = job.get("company", "").lower().strip()
                j_title = job.get("title", "").lower().strip()
                if j_comp == company_clean and j_title == title_clean:
                    return job.get("description", ""), job.get("matchedSkills", []), job.get("missingSkills", [])
                    
            for job in jobs:
                j_comp = job.get("company", "").lower().strip()
                j_title = job.get("title", "").lower().strip()
                if (company_clean in j_comp or j_comp in company_clean) and (title_clean in j_title or j_title in title_clean):
                    return job.get("description", ""), job.get("matchedSkills", []), job.get("missingSkills", [])
                    
            for job in jobs:
                j_comp = job.get("company", "").lower().strip()
                if company_clean in j_comp or j_comp in company_clean:
                    return job.get("description", ""), job.get("matchedSkills", []), job.get("missingSkills", [])
    except Exception as e:
        print(f"Error loading jobs.json: {e}")
    return "", [], []

def align_bullet_text(text, opt_data):
    lang_str = opt_data["lang_str"]
    backend_str = opt_data["backend_str"]
    cloud_str = opt_data["cloud_str"]
    container_str = opt_data["container_str"]
    data_str = opt_data["data_str"]
    obs_str = opt_data["obs_str"]
    cicd_str = opt_data["cicd_str"]
    
    # Custom derivations or fallbacks
    iac_str = opt_data.get("iac_str") or (f"{cloud_str} and Terraform" if "terraform" in cloud_str.lower() else "Terraform and Ansible")
    api_str = opt_data.get("api_str") or (f"{backend_str} REST APIs" if "fastapi" in backend_str.lower() else "RESTful APIs")
    security_str = opt_data.get("security_str") or "OAuth2/JWT"
    
    patterns = [
        (r'\basync Python 3\.12\b', f"async {lang_str}"),
        (r'\bPython 3\.12\b', lang_str),
        (r'\bPython\b', lang_str),
        
        (r'\bmulti-account AWS\b', f"multi-account {cloud_str}"),
        (r'\bAmazon EKS\b', f"{container_str} (EKS)"),
        (r'\bAWS\b', cloud_str),
        
        (r'\bTerraform and AWS CDK \(Python\)(?![a-zA-Z0-9])', iac_str),
        (r'\bTerraform and Ansible\b', iac_str),
        (r'\bTerraform\b', iac_str),
        
        (r'\bJenkins, GitHub Actions, and ArgoCD\b', cicd_str),
        (r'\bGitHub Actions CI/CD\b', f"{cicd_str} CI/CD"),
        (r'\bGitHub Actions\b', cicd_str),
        (r'\bJenkins CI/CD\b', f"{cicd_str} CI/CD"),
        
        (r'\bApache Kafka \(Strimzi on Kubernetes\)(?![a-zA-Z0-9])', f"{data_str} on Kubernetes"),
        (r'\bApache Kafka\b', data_str),
        (r'\bconfluent-kafka-python\b', f"{data_str} drivers"),
        (r'\bMongoDB \(PyMongo/Motor\)(?![a-zA-Z0-9])', data_str),
        (r'\bMongoDB\b', data_str),
        (r'\bPostgreSQL\b', data_str),
        (r'\bMS SQL Server\b', data_str),
        
        (r'\bCelery \+ Redis\b', f"{backend_str} async tasks"),
        (r'\bRedis\b', data_str),
        
        (r'\bFastAPI, SQLAlchemy\b', backend_str),
        (r'\bFastAPI, and asyncio\b', backend_str),
        (r'\bFastAPI/Flask\b', backend_str),
        (r'\bFastAPI\b', backend_str),
        (r'\bDjango, Flask, and SQLAlchemy ORM\b', backend_str),
        (r'\bDjango\b', backend_str),
        (r'\bFlask\b', backend_str),
        
        (r'\bSplunk, Prometheus, and Grafana\b', obs_str),
        (r'\bGrafana\b', obs_str),
        (r'\bCloudWatch\b', obs_str),
        
        (r'\bOAuth2/JWT\b', security_str),
        (r'\bOAuth2\b', security_str),
        
        (r'\bpyodbc/SQLAlchemy\b', 'database connectors'),
        (r'\bzeep\b', 'integration libraries')
    ]
    
    # Sort patterns by pattern length descending to match most specific longest pattern first
    patterns.sort(key=lambda x: len(x[0]), reverse=True)
    
    # Build combined pattern
    parts = [f"({pat})" for pat, _ in patterns]
    combined_pattern = re.compile('|'.join(parts), flags=re.IGNORECASE)
    
    def repl_cb(match):
        for idx, (_, repl) in enumerate(patterns):
            if match.group(idx + 1) is not None:
                return repl
        return match.group(0)
        
    return combined_pattern.sub(repl_cb, text)


def get_experience_bullets(doc):
    company_bullets = {"hsbc": [], "saluki": [], "scotia": []}
    current_company = None
    found_experience = False
    
    for p in doc.paragraphs:
        text_clean = p.text.lower().strip()
        if not found_experience:
            if text_clean == "experience":
                found_experience = True
            continue
            
        style_name = p.style.name.lower()
        is_bullet = "list" in style_name or "bullet" in style_name or p.text.strip().startswith(('\u2022', '\u25e6', '\u25aa', '-', '*'))
        
        if not is_bullet and p.text.strip():
            if "hsbc" in text_clean:
                current_company = "hsbc"
            elif "saluki" in text_clean:
                current_company = "saluki"
            elif "scotia" in text_clean or "client-scotia" in text_clean:
                current_company = "scotia"
            elif "education" in text_clean or "technical skills" in text_clean:
                current_company = None
        elif is_bullet and current_company:
            company_bullets[current_company].append(p)
            
    return company_bullets

def split_by_comma_outside_parentheses(text):
    parts = []
    current = []
    paren_depth = 0
    for char in text:
        if char == '(':
            paren_depth += 1
            current.append(char)
        elif char == ')':
            paren_depth -= 1
            current.append(char)
        elif char == ',' and paren_depth == 0:
            parts.append("".join(current).strip())
            current = []
        else:
            current.append(char)
    if current:
        parts.append("".join(current).strip())
    return [p for p in parts if p]

def get_subdivision_index(skill, category):
    s_lower = skill.lower().strip()
    if category == "languages":
        langs = ["python", "javascript", "typescript", "sql", "pl/sql", "bash", "yaml", "groovy", "java", "c++", "golang", "go", "ruby", "c#", "scala", "c", "rust", "html", "css", "perl", "powershell", "shell"]
        if any(x == s_lower for x in langs) or s_lower.startswith("c++") or s_lower.startswith("python"):
            return 0
        return 1
    elif category == "backend":
        protocols = ["rest", "restful", "graphql", "grpc", "soap", "xml", "openapi", "swagger", "api design", "soap/xml", "apis", "api"]
        concepts = ["microservices", "event-driven architecture", "distributed systems", "systems design", "microservice", "event-driven", "concurrency", "asynchronous"]
        if any(x in s_lower for x in protocols):
            return 1
        if any(x in s_lower for x in concepts):
            return 2
        return 0
    elif category == "cloud":
        tools = ["terraform", "cdk", "pulumi", "cloudformation", "ansible", "packer", "vagrant", "chef", "puppet", "iac"]
        concepts = ["iac", "multi-account", "hybrid", "cloud computing", "architecture"]
        if any(x in s_lower for x in tools):
            return 1
        if any(x in s_lower for x in concepts):
            return 2
        return 0
    elif category == "containers":
        containers = ["docker", "kubernetes", "helm", "istio", "k8s", "kube", "container", "containerization", "mesh"]
        cicd = ["jenkins", "actions", "gitlab", "argocd", "pipeline", "pipelines", "travis", "circleci", "codepipeline"]
        if any(x in s_lower for x in containers):
            return 0
        if any(x in s_lower for x in cicd):
            return 1
        return 2
    elif category == "data":
        brokers = ["kafka", "rabbitmq", "sqs", "sns", "activemq", "nats", "message", "broker", "event streaming", "streaming"]
        concepts = ["optimization", "indexing", "modeling", "caching", "query", "tuning", "replication", "sharding"]
        if any(x in s_lower for x in brokers):
            return 1
        if any(x in s_lower for x in concepts):
            return 2
        return 0
    elif category == "testing":
        obs_tools = ["prometheus", "grafana", "elk", "datadog", "splunk", "cloudwatch", "opentelemetry", "dynatrace", "appdynamics", "tempo", "jaeger"]
        obs_concepts = ["logging", "alerting", "mttr", "monitoring", "telemetry", "observability", "metrics", "tracing"]
        practices = ["tdd", "code review", "code reviews", "peer review"]
        concepts = ["unit", "integration", "load", "e2e", "functional", "regression", "performance", "testing"]
        if any(x in s_lower for x in obs_tools):
            return 3
        if any(x in s_lower for x in obs_concepts):
            return 4
        if any(x in s_lower for x in practices):
            return 2
        if any(x in s_lower for x in concepts):
            return 1
        return 0
    elif category == "security":
        quality = ["sonarqube", "pylint", "mypy", "trivy", "snyk", "black", "flake8", "lint", "linter", "static analysis"]
        practices = ["agile", "scrum", "jira", "confluence", "git", "github", "gitlab", "linux", "unix", "support", "on-call"]
        if any(x in s_lower for x in quality):
            return 1
        if any(x in s_lower for x in practices):
            return 2
        return 0
    return 0

def reorder_skills_paragraph(text, prefix, target_skills):
    if not text.strip().startswith(prefix):
        return None, None
    content = text[len(prefix):].strip()
    
    # Normalize by replacing any pipe character with comma
    content = content.replace(" | ", ", ").replace("|", ", ")
    
    items = split_by_comma_outside_parentheses(content)
    
    matching_items = []
    other_items = []
    seen_items = set()
    
    for item in items:
        item_cap = capitalize_skill(item)
        item_lower = item_cap.lower()
        if item_lower in seen_items:
            continue
        seen_items.add(item_lower)
        
        is_match = False
        for ts in target_skills:
            ts_clean = ts.lower().strip()
            if ts_clean == item_lower or f" {ts_clean} " in f" {item_lower} " or item_lower.startswith(ts_clean + " ") or item_lower.endswith(" " + ts_clean):
                is_match = True
                break
        if is_match:
            matching_items.append(item_cap)
        else:
            other_items.append(item_cap)
            
    reordered_items = matching_items + other_items
    return prefix, ", ".join(reordered_items)

def main():
    if len(sys.argv) < 6:
        print("Usage: python modify_resume.py <base_path> <output_path> <company> <title> <keywords>")
        sys.exit(1)
        
    base_path = sys.argv[1]
    output_path = sys.argv[2]
    company = sys.argv[3]
    title = sys.argv[4]
    keywords = sys.argv[5]
    
    # Exclude processing if base_path is the old template name (force new template path)
    if "Sai_Swaroop_Reddy_Resume_.docx" in base_path:
        base_path = base_path.replace("Sai_Swaroop_Reddy_Resume_.docx", "Sai_Swaroop_Reddy_Resume_ATS.docx")
        
    if not os.path.exists(base_path):
        print(f"Error: Base resume not found at {base_path}")
        sys.exit(1)
        
    # Ensure parent output directory exists
    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)
        
    # Copy base to output
    shutil.copy(base_path, output_path)
    print(f"Copied base resume to output path: {output_path}")
    
    # Load document
    doc = Document(output_path)
    
    # Parse missing keywords from CLI
    skill_list = [s.strip() for s in keywords.split(",") if s.strip() and s.strip() != "General Optimization"]
    
    if len(sys.argv) > 6:
        matched_keywords_str = sys.argv[6]
        matched_skills = [s.strip() for s in matched_keywords_str.split(",") if s.strip()]
    else:
        full_text_scan = " ".join([p.text for p in doc.paragraphs])
        matched_skills = []
        for sk in CATEGORIES_MAPPING.keys():
            if is_skill_present(full_text_scan, sk):
                matched_skills.append(sk)
            
    selected_llm = sys.argv[7] if len(sys.argv) > 7 else "claude"
    
    # --- Tech Stack Integration & Scrapes ---
    # 1. DB (jobs.json) Lookup
    jd_text, db_matched, db_missing = get_job_description_from_db(company, title)
    
    # 2. DDG Scrape
    print(f"Searching DuckDuckGo for {company} {title} tech stack...")
    ddg_skills = search_ddg_tech_stack(company, title)
    print(f"DuckDuckGo found stack skills: {ddg_skills}")
    
    # 3. Extract skills from JD text
    jd_skills = []
    if jd_text:
        jd_text_lower = jd_text.lower()
        for sk in SKILL_CAPITALIZATION.keys():
            escaped_pat = re.escape(sk)
            regex_pattern = f"(?<![a-z0-9+#]){escaped_pat}(?![a-z0-9+#])"
            if re.search(regex_pattern, jd_text_lower):
                jd_skills.append(sk)
                
    # Combine target skills
    all_target_skills = list(set(
        [s.lower().strip() for s in skill_list] + 
        [s.lower().strip() for s in matched_skills] + 
        [s.lower().strip() for s in db_matched] + 
        [s.lower().strip() for s in db_missing] + 
        jd_skills + 
        ddg_skills
    ))
    all_target_skills = [s for s in all_target_skills if s]
    
    # Capitalize professional skill strings for skills paragraphs
    target_matched = [s for s in all_target_skills if s in SAI_BASE_SKILLS]
    target_missing = [s for s in all_target_skills if s not in SAI_BASE_SKILLS]
    
    capitalized_target_matched = [capitalize_skill(s) for s in target_matched]
    capitalized_target_missing = [capitalize_skill(s) for s in target_missing]
    
    # Generate opt_data mapping strings for bullets
    opt_data = generate_optimized_bullets(company, title, target_matched, target_missing, selected_llm)
    
    # --- Update Subtitle ---
    # Find primary languages, cloud, and container systems of target
    primary_langs = get_bullet_skills([s for s in target_matched if get_category(s) == "languages"], 2, "Python")
    primary_cloud = get_bullet_skills([s for s in target_matched if get_category(s) == "cloud"], 1, "AWS")
    primary_containers = get_bullet_skills([s for s in target_matched if get_category(s) == "containers"], 1, "Kubernetes")
    
    new_subtitle = f"Software Engineer | Backend & Cloud | {primary_langs} | {primary_cloud} | {primary_containers}"
    for p in doc.paragraphs:
        if "software engineer" in p.text.lower() and ("backend" in p.text.lower() or "cloud" in p.text.lower()):
            replace_paragraph_text(p, new_subtitle)
            print(f"Updated Subtitle: {new_subtitle}")
            break

    # --- Update Summary ---
    lang_str_primary = get_bullet_skills([s for s in target_matched if get_category(s) == "languages"], 2, "Python")
    backend_str_primary = get_bullet_skills([s for s in target_matched if get_category(s) == "backend"], 2, "FastAPI")
    cloud_str_primary = get_bullet_skills([s for s in target_matched if get_category(s) == "cloud"], 2, "AWS and Terraform")
    container_str_primary = get_bullet_skills([s for s in target_matched if get_category(s) == "containers"], 2, "Kubernetes (EKS) and Docker")
    data_str_primary = get_bullet_skills([s for s in target_matched if get_category(s) == "data"], 2, "Apache Kafka and PostgreSQL")
    api_str_primary = get_bullet_skills([s for s in target_matched if get_category(s) in ["backend", "languages"] and s in ["rest", "restful", "graphql", "grpc", "soap", "xml", "openapi", "swagger", "fastapi", "django", "flask", "spring", "spring boot"]], 2, "RESTful APIs")
    
    new_summary = (
        f"Software engineer with 4+ years of experience designing and scaling production backend systems and cloud-native platforms on {cloud_str_primary} to drive technical excellence at {company} as a {title}. "
        f"Currently developing high-throughput {lang_str_primary} microservices and backend services on {container_str_primary} for HSBC's core platform, processing 1M+ daily transactions with sub-second latency. "
        f"Proficient in distributed systems architecture, designing secure APIs using {api_str_primary}, provisioning infrastructure via {cloud_str_primary}, and managing data pipelines with {data_str_primary} under GitOps continuous delivery. "
        "AWS Certified Developer + Solutions Architect Associate."
    )
    
    found_summary = False
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if p_text.lower() == "summary":
            found_summary = True
            continue
        if found_summary and p_text:
            replace_paragraph_text(p, new_summary)
            print("Updated Summary paragraph dynamically.")
            break

    # --- Update experience titles ---
    first_lang = get_bullet_skills([s for s in target_matched if get_category(s) == "languages"], 1, "Python")
    for p in doc.paragraphs:
        if "associate software engineer - python" in p.text.lower():
            replace_paragraph_text(p, p.text.replace("Python", first_lang))
            print("Updated Scotia Bank role title language.")

    # --- Update experience bullet points for all three companies ---
    company_bullets = get_experience_bullets(doc)
    for company_key, paragraphs in company_bullets.items():
        for p in paragraphs:
            aligned_text = align_bullet_text(p.text, opt_data)
            replace_paragraph_text(p, aligned_text)
        print(f"Dynamically aligned {len(paragraphs)} bullet points for {company_key}")

    # --- Update Core Competencies ---
    # Add matched first, then missing
    competency_skills = []
    for s in target_matched + target_missing:
        cap_skill = capitalize_skill(s)
        if cap_skill.lower() == "python":
            competency_skills.append("Python Microservices")
        elif cap_skill.lower() == "fastapi":
            competency_skills.append("RESTful API Design (FastAPI)")
        elif cap_skill.lower() == "aws":
            competency_skills.append("AWS Cloud Architecture")
        elif cap_skill.lower() == "kubernetes":
            competency_skills.append("Kubernetes Orchestration")
        elif cap_skill.lower() == "docker":
            competency_skills.append("Docker Containerization")
        elif cap_skill.lower() in ["ci/cd", "jenkins", "github actions"]:
            competency_skills.append("CI/CD & DevOps GitOps")
        elif cap_skill.lower() == "terraform":
            competency_skills.append("Infrastructure-as-Code (Terraform)")
        elif cap_skill.lower() == "kafka":
            competency_skills.append("Event-Driven Architecture (Kafka)")
        elif cap_skill.lower() == "postgresql":
            competency_skills.append("Relational Databases (PostgreSQL)")
        elif cap_skill.lower() in ["prometheus", "grafana", "splunk", "observability"]:
            competency_skills.append("Telemetry & Observability")
        else:
            if cap_skill not in competency_skills:
                competency_skills.append(cap_skill)
                
    seen_comp = set()
    final_comps = []
    for comp in competency_skills:
        comp_clean = comp.lower().strip()
        if comp_clean not in seen_comp:
            seen_comp.add(comp_clean)
            final_comps.append(comp)
            
    if len(final_comps) < 8:
        standard_comps = ["Python Microservices", "RESTful API Design", "Distributed Systems", "AWS Cloud", "Kubernetes & Docker", "CI/CD & GitOps", "Infrastructure-as-Code", "Apache Kafka", "Event-Driven Architecture", "Observability", "OAuth2/JWT", "Agile/Scrum"]
        for sc in standard_comps:
            if sc.lower() not in seen_comp and len(final_comps) < 12:
                seen_comp.add(sc.lower())
                final_comps.append(sc)
                
    new_comp_text = ", ".join(final_comps[:12])
    
    for p in doc.paragraphs:
        if p.text.strip().startswith("Core Competencies:"):
            update_prefixed_paragraph(p, "Core Competencies:", new_comp_text)
            print(f"Updated Core Competencies: Core Competencies: {new_comp_text}")
            break

    # --- Update Technical Skills subcategories ---
    # First, append target skills that are missing in the category
    for sk in capitalized_target_matched + capitalized_target_missing:
        cat = get_category(sk)
        prefix = PREFIX_MAP[cat]
        p_sk = None
        for p in doc.paragraphs:
            if p.text.strip().startswith(prefix):
                p_sk = p
                break
        if p_sk:
            if not is_skill_present(p_sk.text, sk):
                p_text = p_sk.text.strip()
                p_text = p_text.replace(" | ", ", ").replace("|", ", ")
                clean_text = p_text[len(prefix):].strip()
                if clean_text:
                    new_skills = clean_text + f", {sk}"
                else:
                    new_skills = sk
                update_prefixed_paragraph(p_sk, prefix, new_skills)
                print(f"Added '{sk}' to '{prefix}' paragraph")
                    
    # Next, reorder each technical skills paragraph to list target stack first
    for cat, prefix in PREFIX_MAP.items():
        p_sk = None
        for p in doc.paragraphs:
            if p.text.strip().startswith(prefix):
                p_sk = p
                break
        if p_sk:
            pref, skills_body = reorder_skills_paragraph(p_sk.text, prefix, all_target_skills)
            if pref:
                update_prefixed_paragraph(p_sk, pref, skills_body)
                print(f"Reordered technical skills for prefix '{prefix}'")

    doc.save(output_path)
    print(f"Successfully saved optimized resume to {output_path}")

if __name__ == "__main__":
    main()
